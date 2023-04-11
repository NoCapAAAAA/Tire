from _cffi_backend import string
from django.shortcuts import render
from django.urls import reverse_lazy
from django.views.generic import FormView, UpdateView, DetailView, TemplateView, ListView
from . import models as m
# Для создания отчётов
from django.http import HttpResponse
from docx import *
from io import BytesIO
from datetime import datetime, date
import locale

from organization import forms as f


class DirectorHomeView(TemplateView):
    template_name = 'director/home.html'


class DirectorOrdersView(ListView):
    model = m.OrderStorage
    template_name = 'director/orders.html'
    context_object_name = 'orders'
    paginate_by = 7

    def get_queryset(self):
        return m.OrderStorage.objects.all().order_by('-pk')


class DirectorUsersView(TemplateView):
    template_name = 'director/users.html'


class DirectorStuffView(TemplateView):
    template_name = 'director/stuff.html'


class DirectorClientsView(TemplateView):
    template_name = 'director/clients.html'


class DirectorStatisticsView(TemplateView):
    template_name = 'director/statistics.html'


class ManagerHomeView(TemplateView):
    template_name = 'manager/home.html'


class ManagerOrdersView(TemplateView):
    template_name = 'manager/orders.html'


class ManagerEdit(UpdateView):
    template_name = 'manager/edit_profile_manager.html'
    form_class = f.SettingsProfile
    success_url = reverse_lazy('manager_edit')

    def get_object(self, **kwargs):
        return self.request.user


class PassChange(UpdateView):
    form_class = f.CustomPasswordChangeForm
    template_name = 'manager/pass_change.html'
    success_url = reverse_lazy('home')

    def get_object(self):
        return self.request.user


def TestDocument(request):
    locale.setlocale(
        category=locale.LC_ALL,
        locale="Russian"  # Note: do not use "de_DE" as it doesn't work
    )
    current_datetime = datetime.now()
    str_current_datetime = str(current_datetime)
    document = Document()
    docx_title = "report" + str_current_datetime + ".docx"
    # ---- Cover Letter ----
    document.add_paragraph()
    document.add_paragraph("%s" % date.today().strftime('%B %d, %Y'))

    document.add_paragraph('Отчёт о заказах')
    qs2 = m.OrderStorage.objects.all().order_by('pk')
    qs2count = qs2.count() + 1
    print(qs2count)
    table = document.add_table(rows=qs2count, cols=8)
    table.style = 'Table Grid'
    table.cell(0, 0).text = 'Номер заказа'
    table.cell(0, 1).text = 'Клиент'
    table.cell(0, 2).text = 'Размер шины'
    table.cell(0, 3).text = 'Период хранения'
    table.cell(0, 4).text = 'Адрес сервиса'
    table.cell(0, 5).text = 'Статус заказа'
    table.cell(0, 6).text = 'Стоимость заказа'
    table.cell(0, 7).text = 'Оплачен'

    # Creating a table object

    for order in m.OrderStorage.objects.all().order_by('pk'):

        row = order.pk
        table.cell(row, 0).text = str(order.pk)
        table.cell(row, 1).text = str(order.user)
        table.cell(row, 2).text = str(order.size)
        table.cell(row, 3).text = str(order.period)
        if order.adress == None:
            order.adress = "---"
        table.cell(row, 4).text = str(order.adress)
        table.cell(row, 5).text = str(order.get_status_display())
        table.cell(row, 6).text = str(order.price)
        if order.is_payed == True:
            order.is_payed = "Да"
        else:
            order.is_payed = "Нет"
        table.cell(row, 7).text = str(order.is_payed)

    document.add_page_break()

    # Prepare document for download
    # -----------------------------
    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length
    return response
